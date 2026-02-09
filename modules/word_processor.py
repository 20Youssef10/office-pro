"""
Word Processor Module
Professional document editing with support for DOCX, DOC, ODT, RTF, and TXT files.
Enhanced with 20+ Microsoft Office features
"""

import os
import re
from pathlib import Path
from collections import deque

try:
    from PyQt6.QtWidgets import (
        QWidget,
        QVBoxLayout,
        QHBoxLayout,
        QTextEdit,
        QToolBar,
        QPushButton,
        QComboBox,
        QSpinBox,
        QFileDialog,
        QMessageBox,
        QFontDialog,
        QColorDialog,
        QMenu,
        QInputDialog,
        QLabel,
        QFrame,
        QSplitter,
        QDialog,
        QLineEdit,
        QCheckBox,
        QGridLayout,
        QTableWidget,
        QTableWidgetItem,
        QTabWidget,
        QGroupBox,
        QRadioButton,
        QProgressBar,
        QScrollArea,
        QTextBrowser,
        QFileDialog,
        QApplication,
        QSizePolicy,
    )
    from PyQt6.QtCore import Qt, QSize, QTimer, QDateTime, pyqtSignal
    from PyQt6.QtGui import (
        QTextCharFormat,
        QFont,
        QColor,
        QTextCursor,
        QAction,
        QIcon,
        QTextListFormat,
        QKeySequence,
        QShortcut,
        QPixmap,
        QImage,
    )
    from PyQt6.QtPrintSupport import QPrinter, QPrintDialog, QPrintPreviewDialog

    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        DOCX_AVAILABLE = True
    except ImportError:
        DOCX_AVAILABLE = False

    try:
        import odfpy

        ODF_AVAILABLE = True
    except ImportError:
        ODF_AVAILABLE = False

    try:
        from spellchecker import SpellChecker

        SPELL_CHECK_AVAILABLE = True
    except ImportError:
        SPELL_CHECK_AVAILABLE = False

except ImportError as e:
    print(f"Import error in word_processor: {e}")
    raise


class FindReplaceDialog(QDialog):
    """Find and Replace dialog"""

    def __init__(self, parent=None, editor=None):
        super().__init__(parent)
        self.editor = editor
        self.setWindowTitle("Find and Replace")
        self.setMinimumWidth(400)
        self.setup_ui()

    def setup_ui(self):
        layout = QGridLayout(self)

        # Find
        layout.addWidget(QLabel("Find:"), 0, 0)
        self.find_input = QLineEdit()
        self.find_input.setPlaceholderText("Enter text to find...")
        layout.addWidget(self.find_input, 0, 1, 1, 2)

        # Replace
        layout.addWidget(QLabel("Replace:"), 1, 0)
        self.replace_input = QLineEdit()
        self.replace_input.setPlaceholderText("Enter replacement text...")
        layout.addWidget(self.replace_input, 1, 1, 1, 2)

        # Options
        options_group = QGroupBox("Options")
        options_layout = QVBoxLayout()

        self.case_sensitive = QCheckBox("Case sensitive")
        options_layout.addWidget(self.case_sensitive)

        self.whole_words = QCheckBox("Whole words only")
        options_layout.addWidget(self.whole_words)

        options_group.setLayout(options_layout)
        layout.addWidget(options_group, 2, 0, 1, 3)

        # Buttons
        self.find_btn = QPushButton("Find Next")
        self.find_btn.clicked.connect(self.find_next)
        layout.addWidget(self.find_btn, 3, 0)

        self.replace_btn = QPushButton("Replace")
        self.replace_btn.clicked.connect(self.replace)
        layout.addWidget(self.replace_btn, 3, 1)

        self.replace_all_btn = QPushButton("Replace All")
        self.replace_all_btn.clicked.connect(self.replace_all)
        layout.addWidget(self.replace_all_btn, 3, 2)

        self.close_btn = QPushButton("Close")
        self.close_btn.clicked.connect(self.close)
        layout.addWidget(self.close_btn, 4, 1)

    def find_next(self):
        """Find next occurrence"""
        text = self.find_input.text()
        if not text:
            return

        flags = QTextDocument.FindFlag(0)
        if self.case_sensitive.isChecked():
            flags |= QTextDocument.FindFlag.FindCaseSensitively
        if self.whole_words.isChecked():
            flags |= QTextDocument.FindFlag.FindWholeWords

        if self.editor.find(text, flags):
            self.find_btn.setStyleSheet("background-color: #90EE90;")
            QTimer.singleShot(500, lambda: self.find_btn.setStyleSheet(""))
        else:
            # Start from beginning
            cursor = self.editor.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.Start)
            self.editor.setTextCursor(cursor)
            if self.editor.find(text, flags):
                self.find_btn.setStyleSheet("background-color: #90EE90;")
                QTimer.singleShot(500, lambda: self.find_btn.setStyleSheet(""))
            else:
                QMessageBox.information(self, "Not Found", f"'{text}' not found.")

    def replace(self):
        """Replace current occurrence"""
        cursor = self.editor.textCursor()
        if cursor.hasSelection():
            cursor.insertText(self.replace_input.text())
        self.find_next()

    def replace_all(self):
        """Replace all occurrences"""
        find_text = self.find_input.text()
        replace_text = self.replace_input.text()

        if not find_text:
            return

        # Get all text
        cursor = self.editor.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.Start)
        self.editor.setTextCursor(cursor)

        count = 0
        flags = QTextDocument.FindFlag(0)
        if self.case_sensitive.isChecked():
            flags |= QTextDocument.FindFlag.FindCaseSensitively
        if self.whole_words.isChecked():
            flags |= QTextDocument.FindFlag.FindWholeWords

        while self.editor.find(find_text, flags):
            cursor = self.editor.textCursor()
            cursor.insertText(replace_text)
            count += 1

        QMessageBox.information(self, "Replace All", f"Replaced {count} occurrences.")


class DocumentStatisticsDialog(QDialog):
    """Document statistics dialog"""

    def __init__(self, parent=None, text=""):
        super().__init__(parent)
        self.setWindowTitle("Document Statistics")
        self.setMinimumWidth(350)
        self.setup_ui(text)

    def setup_ui(self, text):
        layout = QVBoxLayout(self)

        # Calculate statistics
        words = len(text.split())
        chars_with_spaces = len(text)
        chars_without_spaces = len(text.replace(" ", "").replace("\n", ""))
        paragraphs = len([p for p in text.split("\n") if p.strip()])
        lines = len(text.split("\n"))

        # Create table
        table = QTableWidget(5, 2)
        table.setHorizontalHeaderLabels(["Statistic", "Value"])
        table.setItem(0, 0, QTableWidgetItem("Words"))
        table.setItem(0, 1, QTableWidgetItem(str(words)))
        table.setItem(1, 0, QTableWidgetItem("Characters (with spaces)"))
        table.setItem(1, 1, QTableWidgetItem(str(chars_with_spaces)))
        table.setItem(2, 0, QTableWidgetItem("Characters (no spaces)"))
        table.setItem(2, 1, QTableWidgetItem(str(chars_without_spaces)))
        table.setItem(3, 0, QTableWidgetItem("Paragraphs"))
        table.setItem(3, 1, QTableWidgetItem(str(paragraphs)))
        table.setItem(4, 0, QTableWidgetItem("Lines"))
        table.setItem(4, 1, QTableWidgetItem(str(lines)))

        table.resizeColumnsToContents()
        layout.addWidget(table)

        # Close button
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(self.close)
        layout.addWidget(close_btn)


class SpellCheckDialog(QDialog):
    """Spell check dialog"""

    def __init__(self, parent=None, editor=None):
        super().__init__(parent)
        self.editor = editor
        self.setWindowTitle("Spell Check")
        self.setMinimumWidth(400)
        self.spell = SpellChecker() if SPELL_CHECK_AVAILABLE else None
        self.misspelled_words = []
        self.current_index = 0

        self.setup_ui()
        self.start_spell_check()

    def setup_ui(self):
        layout = QVBoxLayout(self)

        # Current word
        layout.addWidget(QLabel("Not in dictionary:"))
        self.word_label = QLabel()
        self.word_label.setStyleSheet("font-weight: bold; font-size: 16px; color: red;")
        layout.addWidget(self.word_label)

        # Suggestions
        layout.addWidget(QLabel("Suggestions:"))
        self.suggestions_list = QTableWidget()
        self.suggestions_list.setColumnCount(1)
        self.suggestions_list.setHorizontalHeaderLabels(["Suggestions"])
        self.suggestions_list.setSelectionBehavior(
            QTableWidget.SelectionBehavior.SelectRows
        )
        layout.addWidget(self.suggestions_list)

        # Buttons
        btn_layout = QHBoxLayout()

        self.ignore_btn = QPushButton("Ignore")
        self.ignore_btn.clicked.connect(self.ignore_word)
        btn_layout.addWidget(self.ignore_btn)

        self.ignore_all_btn = QPushButton("Ignore All")
        self.ignore_all_btn.clicked.connect(self.ignore_all)
        btn_layout.addWidget(self.ignore_all_btn)

        self.change_btn = QPushButton("Change")
        self.change_btn.clicked.connect(self.change_word)
        btn_layout.addWidget(self.change_btn)

        self.change_all_btn = QPushButton("Change All")
        self.change_all_btn.clicked.connect(self.change_all_words)
        btn_layout.addWidget(self.change_all_btn)

        layout.addLayout(btn_layout)

        # Close button
        self.close_btn = QPushButton("Close")
        self.close_btn.clicked.connect(self.close)
        layout.addWidget(self.close_btn)

    def start_spell_check(self):
        """Start spell checking"""
        if not SPELL_CHECK_AVAILABLE:
            QMessageBox.warning(
                self, "Error", "Spell checker not available. Install pyspellchecker."
            )
            self.close()
            return

        text = self.editor.toPlainText()
        words = re.findall(r"\b[a-zA-Z]+\b", text.lower())
        self.misspelled_words = list(self.spell.unknown(words))

        if not self.misspelled_words:
            QMessageBox.information(
                self, "Spell Check Complete", "No spelling errors found!"
            )
            self.close()
            return

        self.current_index = 0
        self.show_current_word()

    def show_current_word(self):
        """Show current misspelled word"""
        if self.current_index >= len(self.misspelled_words):
            QMessageBox.information(
                self, "Spell Check Complete", "Spell check finished!"
            )
            self.close()
            return

        word = self.misspelled_words[self.current_index]
        self.word_label.setText(word)

        # Get suggestions
        suggestions = list(self.spell.candidates(word))[:10]
        self.suggestions_list.setRowCount(len(suggestions))

        for i, suggestion in enumerate(suggestions):
            self.suggestions_list.setItem(i, 0, QTableWidgetItem(suggestion))

        # Highlight word in editor
        self.highlight_word(word)

    def highlight_word(self, word):
        """Highlight word in editor"""
        cursor = self.editor.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.Start)
        self.editor.setTextCursor(cursor)

        # Find and select the word
        if self.editor.find(word):
            cursor = self.editor.textCursor()
            fmt = QTextCharFormat()
            fmt.setBackground(QColor("yellow"))
            cursor.mergeCharFormat(fmt)

    def ignore_word(self):
        """Ignore current word"""
        self.current_index += 1
        self.show_current_word()

    def ignore_all(self):
        """Ignore all occurrences"""
        current_word = self.misspelled_words[self.current_index]
        self.misspelled_words = [w for w in self.misspelled_words if w != current_word]
        self.show_current_word()

    def change_word(self):
        """Change current word"""
        current = self.suggestions_list.currentItem()
        if current:
            new_word = current.text()
            old_word = self.misspelled_words[self.current_index]

            # Replace in editor
            cursor = self.editor.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.Start)
            self.editor.setTextCursor(cursor)

            if self.editor.find(old_word):
                cursor = self.editor.textCursor()
                cursor.insertText(new_word)

        self.current_index += 1
        self.show_current_word()

    def change_all_words(self):
        """Change all occurrences"""
        current = self.suggestions_list.currentItem()
        if current:
            new_word = current.text()
            old_word = self.misspelled_words[self.current_index]

            # Replace all in editor
            text = self.editor.toPlainText()
            text = text.replace(old_word, new_word)
            self.editor.setPlainText(text)

            # Remove from list
            self.misspelled_words = [w for w in self.misspelled_words if w != old_word]

        self.show_current_word()


class InsertTableDialog(QDialog):
    """Insert table dialog"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Insert Table")
        self.setup_ui()

    def setup_ui(self):
        layout = QGridLayout(self)

        layout.addWidget(QLabel("Rows:"), 0, 0)
        self.rows_spin = QSpinBox()
        self.rows_spin.setRange(1, 100)
        self.rows_spin.setValue(3)
        layout.addWidget(self.rows_spin, 0, 1)

        layout.addWidget(QLabel("Columns:"), 1, 0)
        self.cols_spin = QSpinBox()
        self.cols_spin.setRange(1, 100)
        self.cols_spin.setValue(3)
        layout.addWidget(self.cols_spin, 1, 1)

        layout.addWidget(QLabel("Border Width:"), 2, 0)
        self.border_spin = QSpinBox()
        self.border_spin.setRange(0, 10)
        self.border_spin.setValue(1)
        layout.addWidget(self.border_spin, 2, 1)

        # Buttons
        btn_layout = QHBoxLayout()
        self.ok_btn = QPushButton("Insert")
        self.ok_btn.clicked.connect(self.accept)
        btn_layout.addWidget(self.ok_btn)

        self.cancel_btn = QPushButton("Cancel")
        self.cancel_btn.clicked.connect(self.reject)
        btn_layout.addWidget(self.cancel_btn)

        layout.addLayout(btn_layout, 3, 0, 1, 2)

    def get_table_data(self):
        """Get table parameters"""
        return {
            "rows": self.rows_spin.value(),
            "cols": self.cols_spin.value(),
            "border": self.border_spin.value(),
        }


class WordProcessor(QWidget):
    """Professional Word Processor with rich text editing capabilities and Office features"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.current_file = None
        self.document = None
        self.is_modified = False

        # Undo/Redo system
        self.undo_stack = deque(maxlen=50)
        self.redo_stack = deque(maxlen=50)
        self.last_text = ""
        self.is_undoing = False

        # Format painter
        self.format_painter_active = False
        self.saved_format = None

        # Auto-save
        self.autosave_timer = QTimer()
        self.autosave_timer.timeout.connect(self.autosave)
        self.autosave_enabled = True
        self.autosave_interval = 300000  # 5 minutes

        self.init_ui()
        self.setup_toolbar()
        self.setup_shortcuts()
        self.new_document()

        # Start auto-save
        if self.autosave_enabled:
            self.autosave_timer.start(self.autosave_interval)

    def init_ui(self):
        """Initialize the user interface"""
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # Main toolbar
        self.toolbar = QToolBar()
        self.toolbar.setMovable(False)
        self.toolbar.setStyleSheet("""
            QToolBar {
                background-color: #2b579a;
                border: none;
                padding: 5px;
                spacing: 5px;
            }
            QToolBar QPushButton {
                background-color: rgba(255,255,255,0.1);
                color: white;
                border: none;
                padding: 8px 15px;
                border-radius: 4px;
                font-size: 12px;
            }
            QToolBar QPushButton:hover {
                background-color: rgba(255,255,255,0.2);
            }
            QToolBar QPushButton:pressed {
                background-color: rgba(255,255,255,0.3);
            }
            QToolBar QComboBox {
                background-color: white;
                border: none;
                padding: 5px;
                border-radius: 3px;
                min-width: 120px;
            }
        """)
        layout.addWidget(self.toolbar)

        # Text editor
        self.editor = QTextEdit()
        self.editor.setAcceptRichText(True)
        self.editor.textChanged.connect(self.on_text_changed)
        self.editor.setStyleSheet("""
            QTextEdit {
                background-color: white;
                border: 1px solid #ddd;
                padding: 40px;
                font-family: 'Calibri', 'Arial', sans-serif;
                font-size: 11pt;
            }
        """)

        # Set default document margins
        self.editor.document().setDocumentMargin(72)
        layout.addWidget(self.editor)

        # Status bar
        self.status_frame = QFrame()
        self.status_frame.setStyleSheet("""
            QFrame {
                background-color: #f0f0f0;
                border-top: 1px solid #ddd;
                padding: 5px;
            }
        """)
        status_layout = QHBoxLayout(self.status_frame)
        status_layout.setContentsMargins(10, 5, 10, 5)

        self.page_label = QLabel("Page 1 of 1")
        self.word_count_label = QLabel("Words: 0")
        self.char_count_label = QLabel("Characters: 0")
        self.autosave_label = QLabel("Auto-save: ON")

        status_layout.addWidget(self.page_label)
        status_layout.addStretch()
        status_layout.addWidget(self.word_count_label)
        status_layout.addWidget(self.char_count_label)
        status_layout.addWidget(self.autosave_label)

        layout.addWidget(self.status_frame)

        # Connect text change for undo tracking
        self.editor.textChanged.connect(self.track_changes)

    def setup_toolbar(self):
        """Setup the formatting toolbar with Office features"""
        # File operations
        self.new_btn = QPushButton("New")
        self.new_btn.clicked.connect(self.new_document)
        self.toolbar.addWidget(self.new_btn)

        self.open_btn = QPushButton("Open")
        self.open_btn.clicked.connect(self.open_file_dialog)
        self.toolbar.addWidget(self.open_btn)

        self.save_btn = QPushButton("Save")
        self.save_btn.clicked.connect(self.save)
        self.toolbar.addWidget(self.save_btn)

        self.print_btn = QPushButton("Print")
        self.print_btn.clicked.connect(self.print_document)
        self.toolbar.addWidget(self.print_btn)

        self.toolbar.addSeparator()

        # Undo/Redo
        self.undo_btn = QPushButton("Undo")
        self.undo_btn.clicked.connect(self.undo)
        self.undo_btn.setEnabled(False)
        self.toolbar.addWidget(self.undo_btn)

        self.redo_btn = QPushButton("Redo")
        self.redo_btn.clicked.connect(self.redo)
        self.redo_btn.setEnabled(False)
        self.toolbar.addWidget(self.redo_btn)

        self.toolbar.addSeparator()

        # Font family
        self.font_combo = QComboBox()
        self.font_combo.addItems(
            [
                "Calibri",
                "Arial",
                "Times New Roman",
                "Helvetica",
                "Georgia",
                "Verdana",
                "Courier New",
                "Consolas",
                "Segoe UI",
            ]
        )
        self.font_combo.currentTextChanged.connect(self.change_font_family)
        self.toolbar.addWidget(self.font_combo)

        # Font size
        self.size_combo = QComboBox()
        self.size_combo.addItems([str(i) for i in range(8, 73, 2)])
        self.size_combo.setCurrentText("11")
        self.size_combo.currentTextChanged.connect(self.change_font_size)
        self.toolbar.addWidget(self.size_combo)

        self.toolbar.addSeparator()

        # Formatting buttons
        self.bold_btn = QPushButton("B")
        self.bold_btn.setCheckable(True)
        self.bold_btn.setStyleSheet("font-weight: bold; min-width: 30px;")
        self.bold_btn.clicked.connect(self.toggle_bold)
        self.toolbar.addWidget(self.bold_btn)

        self.italic_btn = QPushButton("I")
        self.italic_btn.setCheckable(True)
        self.italic_btn.setStyleSheet("font-style: italic; min-width: 30px;")
        self.italic_btn.clicked.connect(self.toggle_italic)
        self.toolbar.addWidget(self.italic_btn)

        self.underline_btn = QPushButton("U")
        self.underline_btn.setCheckable(True)
        self.underline_btn.setStyleSheet("text-decoration: underline; min-width: 30px;")
        self.underline_btn.clicked.connect(self.toggle_underline)
        self.toolbar.addWidget(self.underline_btn)

        self.toolbar.addSeparator()

        # Alignment
        self.align_left_btn = QPushButton("Left")
        self.align_left_btn.clicked.connect(
            lambda: self.set_alignment(Qt.AlignmentFlag.AlignLeft)
        )
        self.toolbar.addWidget(self.align_left_btn)

        self.align_center_btn = QPushButton("Center")
        self.align_center_btn.clicked.connect(
            lambda: self.set_alignment(Qt.AlignmentFlag.AlignCenter)
        )
        self.toolbar.addWidget(self.align_center_btn)

        self.align_right_btn = QPushButton("Right")
        self.align_right_btn.clicked.connect(
            lambda: self.set_alignment(Qt.AlignmentFlag.AlignRight)
        )
        self.toolbar.addWidget(self.align_right_btn)

        self.toolbar.addSeparator()

        # Color
        self.color_btn = QPushButton("Color")
        self.color_btn.clicked.connect(self.change_text_color)
        self.toolbar.addWidget(self.color_btn)

        self.highlight_btn = QPushButton("Highlight")
        self.highlight_btn.clicked.connect(self.highlight_text)
        self.toolbar.addWidget(self.highlight_btn)

        self.toolbar.addSeparator()

        # Lists
        self.bullet_btn = QPushButton("• Bullet")
        self.bullet_btn.clicked.connect(self.insert_bullet_list)
        self.toolbar.addWidget(self.bullet_btn)

        self.number_btn = QPushButton("1. Number")
        self.number_btn.clicked.connect(self.insert_numbered_list)
        self.toolbar.addWidget(self.number_btn)

        self.toolbar.addSeparator()

        # Advanced features
        self.find_btn = QPushButton("Find")
        self.find_btn.clicked.connect(self.show_find_replace)
        self.toolbar.addWidget(self.find_btn)

        self.spell_btn = QPushButton("Spell Check")
        self.spell_btn.clicked.connect(self.show_spell_check)
        self.toolbar.addWidget(self.spell_btn)

        self.stats_btn = QPushButton("Statistics")
        self.stats_btn.clicked.connect(self.show_statistics)
        self.toolbar.addWidget(self.stats_btn)

        self.format_painter_btn = QPushButton("Format Painter")
        self.format_painter_btn.setCheckable(True)
        self.format_painter_btn.clicked.connect(self.toggle_format_painter)
        self.toolbar.addWidget(self.format_painter_btn)

        self.image_btn = QPushButton("Insert Image")
        self.image_btn.clicked.connect(self.insert_image)
        self.toolbar.addWidget(self.image_btn)

        self.table_btn = QPushButton("Insert Table")
        self.table_btn.clicked.connect(self.insert_table)
        self.toolbar.addWidget(self.table_btn)

        self.toolbar.addSeparator()

        # Home button
        self.home_btn = QPushButton("← Home")
        self.home_btn.clicked.connect(self.go_home)
        self.toolbar.addWidget(self.home_btn)

    def setup_shortcuts(self):
        """Setup keyboard shortcuts"""
        # Find
        find_shortcut = QShortcut(QKeySequence("Ctrl+F"), self)
        find_shortcut.activated.connect(self.show_find_replace)

        # Print
        print_shortcut = QShortcut(QKeySequence("Ctrl+P"), self)
        print_shortcut.activated.connect(self.print_document)

        # Undo/Redo
        undo_shortcut = QShortcut(QKeySequence("Ctrl+Z"), self)
        undo_shortcut.activated.connect(self.undo)

        redo_shortcut = QShortcut(QKeySequence("Ctrl+Y"), self)
        redo_shortcut.activated.connect(self.redo)

        # Select All
        select_all_shortcut = QShortcut(QKeySequence("Ctrl+A"), self)
        select_all_shortcut.activated.connect(self.editor.selectAll)

    def new_document(self):
        """Create a new document"""
        if self.check_save():
            self.save_state()  # Save state for undo
            self.editor.clear()
            self.current_file = None
            self.is_modified = False
            self.undo_stack.clear()
            self.redo_stack.clear()
            self.update_undo_redo_buttons()
            self.set_default_formatting()
            self.update_title()

    def set_default_formatting(self):
        """Set default document formatting"""
        self.editor.setFont(QFont("Calibri", 11))

    # ========== NEW FEATURES ==========

    def track_changes(self):
        """Track changes for undo/redo"""
        if not self.is_undoing and self.editor.toPlainText() != self.last_text:
            self.save_state()
            self.last_text = self.editor.toPlainText()

    def save_state(self):
        """Save current state to undo stack"""
        if len(self.undo_stack) >= 50:
            self.undo_stack.popleft()
        self.undo_stack.append(self.editor.toHtml())
        self.update_undo_redo_buttons()

    def undo(self):
        """Undo last action"""
        if self.undo_stack:
            self.is_undoing = True
            # Save current state to redo stack
            if len(self.redo_stack) >= 50:
                self.redo_stack.popleft()
            self.redo_stack.append(self.editor.toHtml())

            # Restore previous state
            state = self.undo_stack.pop()
            self.editor.setHtml(state)
            self.is_undoing = False
            self.is_modified = True
            self.update_title()
            self.update_undo_redo_buttons()

    def redo(self):
        """Redo last undone action"""
        if self.redo_stack:
            self.is_undoing = True
            # Save current state to undo stack
            if len(self.undo_stack) >= 50:
                self.undo_stack.popleft()
            self.undo_stack.append(self.editor.toHtml())

            # Restore next state
            state = self.redo_stack.pop()
            self.editor.setHtml(state)
            self.is_undoing = False
            self.is_modified = True
            self.update_title()
            self.update_undo_redo_buttons()

    def update_undo_redo_buttons(self):
        """Update undo/redo button states"""
        self.undo_btn.setEnabled(len(self.undo_stack) > 0)
        self.redo_btn.setEnabled(len(self.redo_stack) > 0)

    def show_find_replace(self):
        """Show find and replace dialog"""
        dialog = FindReplaceDialog(self, self.editor)
        dialog.exec()

    def show_spell_check(self):
        """Show spell check dialog"""
        if not SPELL_CHECK_AVAILABLE:
            QMessageBox.warning(
                self,
                "Spell Check",
                "Spell checker not available. Install pyspellchecker:\n"
                "pip install pyspellchecker",
            )
            return
        dialog = SpellCheckDialog(self, self.editor)
        dialog.exec()

    def show_statistics(self):
        """Show document statistics"""
        dialog = DocumentStatisticsDialog(self, self.editor.toPlainText())
        dialog.exec()

    def toggle_format_painter(self):
        """Toggle format painter"""
        self.format_painter_active = self.format_painter_btn.isChecked()

        if self.format_painter_active:
            # Save current format
            cursor = self.editor.textCursor()
            if cursor.hasSelection():
                self.saved_format = cursor.charFormat()
                self.format_painter_btn.setStyleSheet("background-color: #90EE90;")
            else:
                self.format_painter_active = False
                self.format_painter_btn.setChecked(False)
        else:
            self.saved_format = None
            self.format_painter_btn.setStyleSheet("")

    def apply_format_painter(self):
        """Apply saved format to selection"""
        if self.format_painter_active and self.saved_format:
            cursor = self.editor.textCursor()
            if cursor.hasSelection():
                cursor.setCharFormat(self.saved_format)
                # Optionally turn off after one use
                # self.format_painter_btn.setChecked(False)
                # self.format_painter_active = False

    def insert_image(self):
        """Insert image into document"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Insert Image",
            "",
            "Image Files (*.png *.jpg *.jpeg *.gif *.bmp);;All Files (*)",
        )
        if file_path:
            try:
                image = QImage(file_path)
                if not image.isNull():
                    # Scale image if too large
                    max_width = 600
                    if image.width() > max_width:
                        image = image.scaledToWidth(
                            max_width, Qt.TransformationMode.SmoothTransformation
                        )

                    cursor = self.editor.textCursor()
                    cursor.insertImage(image, file_path)
                    self.is_modified = True
                    self.update_title()
                else:
                    QMessageBox.warning(self, "Error", "Could not load image.")
            except Exception as e:
                QMessageBox.critical(
                    self, "Error", f"Failed to insert image:\n{str(e)}"
                )

    def insert_table(self):
        """Insert table into document"""
        dialog = InsertTableDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            data = dialog.get_table_data()
            self.create_html_table(data["rows"], data["cols"], data["border"])

    def create_html_table(self, rows, cols, border):
        """Create HTML table and insert"""
        html = f'<table border="{border}" cellpadding="5" cellspacing="0" style="border-collapse: collapse;">'

        for r in range(rows):
            html += "<tr>"
            for c in range(cols):
                html += '<td style="border: 1px solid #000; padding: 8px;">Cell</td>'
            html += "</tr>"

        html += "</table><br>"

        cursor = self.editor.textCursor()
        cursor.insertHtml(html)
        self.is_modified = True
        self.update_title()

    def print_document(self):
        """Print document"""
        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        dialog = QPrintDialog(printer, self)

        if dialog.exec() == QPrintDialog.DialogCode.Accepted:
            self.editor.print(printer)

    def print_preview(self):
        """Show print preview"""
        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        preview = QPrintPreviewDialog(printer, self)
        preview.paintRequested.connect(lambda p: self.editor.print(p))
        preview.exec()

    def autosave(self):
        """Auto-save document"""
        if self.is_modified and self.current_file:
            try:
                self.save_as(self.current_file)
                self.autosave_label.setText("Auto-save: Just now")
                QTimer.singleShot(
                    5000, lambda: self.autosave_label.setText("Auto-save: ON")
                )
            except:
                self.autosave_label.setText("Auto-save: Failed")

    def generate_table_of_contents(self):
        """Generate table of contents from headings"""
        text = self.editor.toPlainText()
        lines = text.split("\n")

        toc = ["Table of Contents", "=" * 40, ""]

        for i, line in enumerate(lines):
            if line.strip().startswith("#"):
                # Count heading level
                level = len(line) - len(line.lstrip("#"))
                title = line.lstrip("#").strip()
                indent = "    " * (level - 1)
                toc.append(f"{indent}{title}")

        return "\n".join(toc)

    def insert_table_of_contents(self):
        """Insert table of contents at cursor"""
        toc = self.generate_table_of_contents()
        cursor = self.editor.textCursor()
        cursor.insertText(toc)

    def insert_page_break(self):
        """Insert page break"""
        cursor = self.editor.textCursor()
        cursor.insertHtml("<div style='page-break-after: always;'></div>")

    def insert_hyperlink(self):
        """Insert hyperlink"""
        url, ok = QInputDialog.getText(self, "Insert Hyperlink", "URL:")
        if ok and url:
            text, ok2 = QInputDialog.getText(
                self, "Link Text", "Display text:", text=url
            )
            if ok2:
                cursor = self.editor.textCursor()
                if cursor.hasSelection():
                    cursor.insertHtml(f'<a href="{url}">{cursor.selectedText()}</a>')
                else:
                    cursor.insertHtml(f'<a href="{url}">{text}</a>')

    def insert_header_footer(self):
        """Insert header or footer"""
        choices = ["Header", "Footer"]
        choice, ok = QInputDialog.getItem(self, "Insert", "Select:", choices, 0, False)
        if ok:
            text, ok2 = QInputDialog.getText(self, f"Insert {choice}", "Text:")
            if ok2:
                cursor = self.editor.textCursor()
                if choice == "Header":
                    cursor.movePosition(QTextCursor.MoveOperation.Start)
                    self.editor.setTextCursor(cursor)
                    cursor.insertHtml(
                        f'<div style="border-bottom: 1px solid #ccc; padding-bottom: 5px; margin-bottom: 20px;"><i>{text}</i></div>'
                    )
                else:
                    cursor.movePosition(QTextCursor.MoveOperation.End)
                    self.editor.setTextCursor(cursor)
                    cursor.insertHtml(
                        f'<div style="border-top: 1px solid #ccc; padding-top: 5px; margin-top: 20px; text-align: center;"><i>{text}</i></div>'
                    )

    def set_line_spacing(self):
        """Set line spacing"""
        options = ["1.0", "1.15", "1.5", "2.0", "2.5", "3.0"]
        spacing, ok = QInputDialog.getItem(
            self, "Line Spacing", "Select spacing:", options, 1, False
        )
        if ok:
            cursor = self.editor.textCursor()
            block_fmt = cursor.blockFormat()
            block_fmt.setLineHeight(int(float(spacing) * 100), 1)
            cursor.setBlockFormat(block_fmt)

    def add_watermark(self):
        """Add watermark to document"""
        text, ok = QInputDialog.getText(self, "Add Watermark", "Watermark text:")
        if ok and text:
            # Insert as background text
            html = f"""
            <div style="position: fixed; top: 50%; left: 50%; transform: translate(-50%, -50%) rotate(-45deg); 
                       font-size: 72px; color: rgba(200, 200, 200, 0.3); pointer-events: none; z-index: -1;">
                {text}
            </div>
            """
            cursor = self.editor.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.Start)
            self.editor.setTextCursor(cursor)
            cursor.insertHtml(html)

    def export_as_pdf(self):
        """Export document as PDF"""
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export as PDF", "", "PDF Files (*.pdf);;All Files (*)"
        )
        if file_path:
            try:
                printer = QPrinter(QPrinter.PrinterMode.HighResolution)
                printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
                printer.setOutputFileName(file_path)
                self.editor.document().print(printer)
                QMessageBox.information(
                    self, "Export Complete", f"Document exported to {file_path}"
                )
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Export failed:\n{str(e)}")

    def open_file_dialog(self):
        """Open file dialog"""
        if self.check_save():
            file_path, _ = QFileDialog.getOpenFileName(
                self,
                "Open Document",
                "",
                "Word Documents (*.docx *.doc);;All Files (*)",
            )
            if file_path:
                self.open_document(file_path)

    def open_document(self, file_path):
        """Open a document file"""
        try:
            self.save_state()  # Save state before opening
            if file_path.endswith(".docx") and DOCX_AVAILABLE:
                self.open_docx(file_path)
            elif file_path.endswith(".txt"):
                with open(file_path, "r", encoding="utf-8") as f:
                    self.editor.setPlainText(f.read())
                self.current_file = file_path
            else:
                # Try to read as plain text
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        self.editor.setPlainText(f.read())
                    self.current_file = file_path
                except:
                    QMessageBox.warning(self, "Error", "Could not open file.")
                    return

            self.is_modified = False
            self.update_title()

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to open document:\n{str(e)}")

    def open_docx(self, file_path):
        """Open a DOCX file"""
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, "Error", "python-docx not installed.")
            return

        doc = Document(file_path)
        self.editor.clear()

        # Convert docx to HTML-like content
        html_content = []
        html_content.append("<html><body>")

        for para in doc.paragraphs:
            if para.text.strip():
                alignment = "left"
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    alignment = "center"
                elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    alignment = "right"

                html_content.append(f'<p align="{alignment}">')

                for run in para.runs:
                    text = run.text.replace("<", "&lt;").replace(">", "&gt;")
                    style = []

                    if run.bold:
                        style.append("font-weight: bold;")
                    if run.italic:
                        style.append("font-style: italic;")
                    if run.underline:
                        style.append("text-decoration: underline;")
                    if run.font.size:
                        style.append(f"font-size: {run.font.size.pt}pt;")
                    if run.font.name:
                        style.append(f"font-family: {run.font.name};")

                    if style:
                        html_content.append(
                            f'<span style="{"".join(style)}">{text}</span>'
                        )
                    else:
                        html_content.append(text)

                html_content.append("</p>")
            else:
                html_content.append("<br>")

        html_content.append("</body></html>")

        self.editor.setHtml("".join(html_content))
        self.current_file = file_path

    def save(self):
        """Save the current document"""
        if self.current_file:
            self.save_as(self.current_file)
        else:
            self.save_as_dialog()

    def save_as_dialog(self):
        """Save as dialog"""
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Document",
            "",
            "Word Document (*.docx);;Text File (*.txt);;HTML (*.html);;PDF (*.pdf);;All Files (*)",
        )
        if file_path:
            self.save_as(file_path)

    def save_as(self, file_path):
        """Save document to file"""
        try:
            if file_path.endswith(".docx") and DOCX_AVAILABLE:
                self.save_as_docx(file_path)
            elif file_path.endswith(".html") or file_path.endswith(".htm"):
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(self.editor.toHtml())
            elif file_path.endswith(".pdf"):
                self.export_as_pdf()
                return
            else:
                # Default to text
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(self.editor.toPlainText())

            self.current_file = file_path
            self.is_modified = False
            self.update_title()

            if self.parent:
                self.parent.status_bar.showMessage(f"Saved: {file_path}", 3000)

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save document:\n{str(e)}")

    def save_as_docx(self, file_path):
        """Save as DOCX format"""
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, "Error", "python-docx not installed.")
            return

        doc = Document()

        # Get plain text and add to document
        text = self.editor.toPlainText()
        paragraphs = text.split("\n")

        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text)
            else:
                doc.add_paragraph()

        doc.save(file_path)

    def check_save(self):
        """Check if document needs saving"""
        if self.is_modified:
            reply = QMessageBox.question(
                self,
                "Unsaved Changes",
                "The document has unsaved changes. Save now?",
                QMessageBox.StandardButton.Save
                | QMessageBox.StandardButton.Discard
                | QMessageBox.StandardButton.Cancel,
            )

            if reply == QMessageBox.StandardButton.Save:
                self.save()
                return True
            elif reply == QMessageBox.StandardButton.Cancel:
                return False

        return True

    def on_text_changed(self):
        """Handle text changes"""
        if not self.is_modified:
            self.is_modified = True
            self.update_title()

    def update_title(self):
        """Update window title"""
        if self.parent:
            filename = "Untitled"
            if self.current_file:
                filename = Path(self.current_file).name
            modified = "*" if self.is_modified else ""
            self.parent.setWindowTitle(
                f"Office Pro - Word Processor - {filename}{modified}"
            )

    def update_status(self):
        """Update status bar information"""
        text = self.editor.toPlainText()
        words = len(text.split())
        chars = len(text)

        self.word_count_label.setText(f"Words: {words}")
        self.char_count_label.setText(f"Characters: {chars}")

    def change_font_family(self, font_name):
        """Change font family"""
        cursor = self.editor.textCursor()
        if cursor.hasSelection():
            fmt = QTextCharFormat()
            fmt.setFontFamily(font_name)
            cursor.mergeCharFormat(fmt)
        else:
            font = self.editor.font()
            font.setFamily(font_name)
            self.editor.setFont(font)

    def change_font_size(self, size):
        """Change font size"""
        try:
            size = int(size)
            cursor = self.editor.textCursor()
            if cursor.hasSelection():
                fmt = QTextCharFormat()
                fmt.setFontPointSize(size)
                cursor.mergeCharFormat(fmt)
            else:
                font = self.editor.font()
                font.setPointSize(size)
                self.editor.setFont(font)
        except ValueError:
            pass

    def toggle_bold(self):
        """Toggle bold formatting"""
        cursor = self.editor.textCursor()
        fmt = QTextCharFormat()
        fmt.setFontWeight(
            QFont.Weight.Bold if self.bold_btn.isChecked() else QFont.Weight.Normal
        )
        cursor.mergeCharFormat(fmt)

    def toggle_italic(self):
        """Toggle italic formatting"""
        cursor = self.editor.textCursor()
        fmt = QTextCharFormat()
        fmt.setFontItalic(self.italic_btn.isChecked())
        cursor.mergeCharFormat(fmt)

    def toggle_underline(self):
        """Toggle underline formatting"""
        cursor = self.editor.textCursor()
        fmt = QTextCharFormat()
        fmt.setFontUnderline(self.underline_btn.isChecked())
        cursor.mergeCharFormat(fmt)

    def set_alignment(self, alignment):
        """Set text alignment"""
        self.editor.setAlignment(alignment)

    def change_text_color(self):
        """Change text color"""
        color = QColorDialog.getColor()
        if color.isValid():
            cursor = self.editor.textCursor()
            fmt = QTextCharFormat()
            fmt.setForeground(color)
            cursor.mergeCharFormat(fmt)

    def highlight_text(self):
        """Highlight selected text"""
        color = QColorDialog.getColor()
        if color.isValid():
            cursor = self.editor.textCursor()
            fmt = QTextCharFormat()
            fmt.setBackground(color)
            cursor.mergeCharFormat(fmt)

    def insert_bullet_list(self):
        """Insert bullet list"""
        cursor = self.editor.textCursor()
        cursor.insertList(QTextListFormat.Style.ListDisc)

    def insert_numbered_list(self):
        """Insert numbered list"""
        cursor = self.editor.textCursor()
        cursor.insertList(QTextListFormat.Style.ListDecimal)

    def go_home(self):
        """Return to main menu"""
        if self.check_save():
            if self.parent:
                self.parent.show_main_menu()

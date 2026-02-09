"""
Word Processor Mobile Module
Kivy-based mobile-optimized word processor
"""

import os
from pathlib import Path

try:
    from kivy.uix.screenmanager import Screen
    from kivy.uix.boxlayout import BoxLayout
    from kivy.uix.textinput import TextInput
    from kivy.uix.button import Button
    from kivy.uix.gridlayout import GridLayout
    from kivy.uix.scrollview import ScrollView
    from kivy.uix.popup import Popup
    from kivy.uix.label import Label
    from kivy.uix.filechooser import FileChooserListView
    from kivy.core.window import Window
    from kivy.properties import ObjectProperty
    from kivy.utils import platform

    if platform == "android":
        from android.permissions import request_permissions, Permission
        from android.storage import primary_external_storage_path

    KIVY_AVAILABLE = True
except ImportError:
    KIVY_AVAILABLE = False
    print("Kivy not available")

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordProcessorMobile(Screen):
    """Mobile word processor screen"""

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.current_file = None
        self.is_modified = False
        self.build_ui()

    def build_ui(self):
        """Build mobile UI"""
        layout = BoxLayout(orientation="vertical")

        # Toolbar
        toolbar = GridLayout(cols=5, size_hint_y=0.1, spacing=5, padding=5)
        toolbar.add_widget(Button(text="New", on_press=self.new_document))
        toolbar.add_widget(Button(text="Open", on_press=self.open_file))
        toolbar.add_widget(Button(text="Save", on_press=self.save_file))
        toolbar.add_widget(Button(text="Back", on_press=self.go_back))
        layout.add_widget(toolbar)

        # Text editor
        scroll = ScrollView()
        self.editor = TextInput(
            multiline=True,
            font_size="16sp",
            background_color=(1, 1, 1, 1),
            foreground_color=(0, 0, 0, 1),
        )
        scroll.add_widget(self.editor)
        layout.add_widget(scroll)

        # Status bar
        status = BoxLayout(size_hint_y=0.05)
        self.status_label = Label(text="Ready", size_hint_x=0.5, font_size="12sp")
        status.add_widget(self.status_label)
        layout.add_widget(status)

        self.add_widget(layout)

    def new_document(self, instance):
        """Create new document"""
        if self.is_modified:
            self.show_confirm_dialog()
        else:
            self.editor.text = ""
            self.current_file = None
            self.is_modified = False
            self.update_status()

    def open_file(self, instance):
        """Open file dialog"""
        if platform == "android":
            request_permissions(
                [Permission.READ_EXTERNAL_STORAGE, Permission.WRITE_EXTERNAL_STORAGE]
            )

        content = BoxLayout(orientation="vertical")
        filechooser = FileChooserListView(
            path=self.get_documents_path(), filters=["*.txt", "*.docx", "*.doc"]
        )
        content.add_widget(filechooser)

        btn_layout = BoxLayout(size_hint_y=0.1)
        btn_layout.add_widget(
            Button(
                text="Open", on_press=lambda x: self.load_file(filechooser.selection)
            )
        )
        btn_layout.add_widget(
            Button(text="Cancel", on_press=lambda x: self.popup.dismiss())
        )
        content.add_widget(btn_layout)

        self.popup = Popup(title="Open File", content=content, size_hint=(0.9, 0.9))
        self.popup.open()

    def get_documents_path(self):
        """Get documents path"""
        if platform == "android":
            return primary_external_storage_path()
        else:
            return os.path.expanduser("~")

    def load_file(self, selection):
        """Load selected file"""
        if selection:
            file_path = selection[0]
            try:
                if file_path.endswith(".docx") and DOCX_AVAILABLE:
                    doc = Document(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    self.editor.text = text
                else:
                    with open(file_path, "r", encoding="utf-8") as f:
                        self.editor.text = f.read()

                self.current_file = file_path
                self.is_modified = False
                self.update_status()
                self.popup.dismiss()

            except Exception as e:
                self.show_error(f"Error loading file: {str(e)}")

    def save_file(self, instance):
        """Save file"""
        if self.current_file:
            self.save_to_file(self.current_file)
        else:
            self.save_as_dialog()

    def save_as_dialog(self):
        """Save as dialog"""
        content = BoxLayout(orientation="vertical")

        filename_input = TextInput(
            hint_text="Enter filename", multiline=False, size_hint_y=0.1
        )
        content.add_widget(filename_input)

        btn_layout = BoxLayout(size_hint_y=0.1)
        btn_layout.add_widget(
            Button(
                text="Save",
                on_press=lambda x: self.perform_save_as(filename_input.text),
            )
        )
        btn_layout.add_widget(
            Button(text="Cancel", on_press=lambda x: self.popup.dismiss())
        )
        content.add_widget(btn_layout)

        self.popup = Popup(title="Save As", content=content, size_hint=(0.9, 0.5))
        self.popup.open()

    def perform_save_as(self, filename):
        """Perform save as"""
        if filename:
            if not filename.endswith(".txt"):
                filename += ".txt"
            file_path = os.path.join(self.get_documents_path(), filename)
            self.save_to_file(file_path)
            self.popup.dismiss()

    def save_to_file(self, file_path):
        """Save text to file"""
        try:
            if file_path.endswith(".docx") and DOCX_AVAILABLE:
                doc = Document()
                for para_text in self.editor.text.split("\n"):
                    doc.add_paragraph(para_text)
                doc.save(file_path)
            else:
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(self.editor.text)

            self.current_file = file_path
            self.is_modified = False
            self.update_status()
            self.show_message("Saved successfully!")

        except Exception as e:
            self.show_error(f"Error saving file: {str(e)}")

    def update_status(self):
        """Update status bar"""
        word_count = len(self.editor.text.split())
        filename = Path(self.current_file).name if self.current_file else "Untitled"
        modified = "*" if self.is_modified else ""
        self.status_label.text = f"{filename}{modified} | Words: {word_count}"

    def show_confirm_dialog(self):
        """Show confirmation dialog"""
        content = BoxLayout(orientation="vertical")
        content.add_widget(Label(text="Save changes before creating new document?"))

        btn_layout = BoxLayout(size_hint_y=0.3)
        btn_layout.add_widget(
            Button(text="Save", on_press=lambda x: self.save_and_new())
        )
        btn_layout.add_widget(
            Button(text="Discard", on_press=lambda x: self.discard_and_new())
        )
        btn_layout.add_widget(
            Button(text="Cancel", on_press=lambda x: self.popup.dismiss())
        )
        content.add_widget(btn_layout)

        self.popup = Popup(
            title="Unsaved Changes", content=content, size_hint=(0.8, 0.4)
        )
        self.popup.open()

    def save_and_new(self):
        """Save and create new"""
        self.save_file(None)
        self.popup.dismiss()
        self.editor.text = ""
        self.current_file = None
        self.is_modified = False

    def discard_and_new(self):
        """Discard and create new"""
        self.popup.dismiss()
        self.editor.text = ""
        self.current_file = None
        self.is_modified = False

    def show_message(self, message):
        """Show message popup"""
        popup = Popup(
            title="Message", content=Label(text=message), size_hint=(0.7, 0.3)
        )
        popup.open()

    def show_error(self, error):
        """Show error popup"""
        popup = Popup(title="Error", content=Label(text=error), size_hint=(0.8, 0.4))
        popup.open()

    def go_back(self, instance):
        """Go back to main menu"""
        self.manager.current = "menu"
